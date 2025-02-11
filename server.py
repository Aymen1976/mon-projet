from flask import Flask, request, send_file
from docx import Document
from io import BytesIO
from reportlab.pdfgen import canvas

app = Flask(__name__)

@app.route('/generate_word', methods=['POST'])
def generate_word():
    data = request.json
    title = data.get('title', 'Titre par défaut')
    content = data.get('content', 'Contenu par défaut')

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    word_file = BytesIO()
    doc.save(word_file)
    word_file.seek(0)

    return send_file(word_file, as_attachment=True, download_name='document.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@app.route('/generate_pdf', methods=['POST'])
def generate_pdf():
    data = request.json
    title = data.get('title', 'Titre par défaut')
    content = data.get('content', 'Contenu par défaut')

    pdf_file = BytesIO()
    pdf = canvas.Canvas(pdf_file)
    pdf.drawString(100, 750, title)
    pdf.drawString(100, 730, content)
    pdf.save()
    pdf_file.seek(0)

    return send_file(pdf_file, as_attachment=True, download_name='document.pdf', mimetype='application/pdf')

if __name__ == '__main__':
    app.run(debug=True)
